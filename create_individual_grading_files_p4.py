# %%
###########################################################################
# This code for generating individual grading sheet from spreadsheet
# required files: spreadsheet file,  the individual grading sheet template
###########################################################################
import os
import sys

import xlrd
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT


# from docx2pdf import convert
# import pdfkit
# from PyPDF2 import PdfFileMerger,PdfFileReader, PdfFileWriter

def LetterGrade(number, total):
    per = number / total
    # if per in Interval(0.95, 1):
    if per > 1.00:
        return 'A+'
    elif per > 0.95:
        return 'A'
    elif per > 0.90 and per <= 0.95:
        return 'A-'
    elif per > 0.85 and per <= 0.90:
        return 'B+'
    elif per > 0.80 and per <= 0.85:
        return 'B'
    elif per > 0.75 and per <= 0.80:
        return 'B-'
    elif per > 0.70 and per <= 0.75:
        return 'C+'
    elif per > 0.65 and per <= 0.70:
        return 'C'
    elif per > 0.60 and per <= 0.65:
        return 'C-'
    elif per > 0.55 and per <= 0.60:
        return 'D+'
    elif per > 0.50 and per <= 0.55:
        return 'D'
    elif per > 0.45 and per <= 0.50:
        return 'D-'
    else:
        return 'F'


def writeIndividualFile(template, firstName, lastName, email, total, q1, q2, q3, q4, q5, q6,
                        comment1, comment2, comment3, comment4, comment5, comment6):
    doc = docx.Document(template)
    table1 = doc.tables[0]  # get the table in the grading sheet

    table1.cell(1, 0).text = firstName + ' ' + lastName  # Student name
    table1.cell(1, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table1.cell(1, 1).text = LetterGrade(total, 189)  # Grade
    table1.cell(1, 1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table1.cell(1, 2).text = str(int(total)) + '/189'  # Points, please remember modify the full score
    table1.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table1.cell(1, 3).text = 'Antu Saha'  # grader name
    table1.cell(1, 3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    table2 = doc.tables[1]

    table2.cell(1, 2).text = str(int(q1))  # question 1 score
    table2.cell(1, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(1, 3).text = comment1  # question 1 comment

    table2.cell(2, 2).text = str(int(q2))
    table2.cell(2, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(2, 3).text = comment2

    table2.cell(3, 2).text = str(int(q3))
    table2.cell(3, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(3, 3).text = comment3

    table2.cell(4, 2).text = str(int(q4))
    table2.cell(4, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(4, 3).text = comment4

    table2.cell(5, 2).text = str(int(q5))
    table2.cell(5, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(5, 3).text = comment5

    table2.cell(6, 2).text = str(int(q6))
    table2.cell(6, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.cell(6, 3).text = comment6

    name = firstName + '_' + lastName + '_' + email  # individual grading sheet file name
    individualFilePath = './outputs/%s.docx' % name
    doc.save(individualFilePath)

    # %%


if __name__ == '__main__':
    outpath = './outputs/'  # the folder stores the individual grading sheets
    if not os.path.exists(outpath):
        os.makedirs(outpath)
    doc_template_path = 'p4_form_f22.docx'  # the template of the individual grading sheet
    excel_path = 'p4_spreadsheet_f22.xls'  # the spreadsheet file
    excel = xlrd.open_workbook(excel_path, encoding_override="utf-8")
    sheet = excel.sheets()[0]
    names = sheet.row_values(1)
    for i in range(len(names)):
        # print(names[i])
        print(str(i) + '  ' + names[i])

    # print(len(names))
    # print(names)
    # ['Last Name', 'First Name', 'User Name', 'Repository', 'Deadline extension',
    #  'Total Points', 'Gitlab vcs', 'Comment', 'Code', 'Comment',
    #  'Junit tests', 'Comment', 'Manual test', 'Comment', 'Bonus Smart Player',
    #  'Comment']

    # 'Notes', 'Grade (total points)',
    #  'Gitlab repository(16)', 'Comments', 'Gitlab issue tracker(24)', 'Comments', 'Class hierarchy(15)', 'Comments',
    #  'Algorithm(40)', 'Comments', 'JUnit tests(82+10)', 'Comments', 'Command line parameter(20)', 'Comments',
    #  'Javadoc(20)', 'Comments', 'Student tests(50, 80)', 'Comments']

    # sys.exit()
    # students row numbers in grading sheet
    start = 2
    end = 67

    for i in range(start, end):  # get values of each record
        list = sheet.row_values(i)
        lastName = list[0]
        firstName = list[1]
        email = list[2]
        # print(list)
        # deadline = list[4]

        # sys.exit()

        total = list[6]

        q1 = list[7]
        comment1 = list[8]
        q2 = list[9]
        comment2 = list[10]
        q3 = list[11]
        comment3 = list[12]
        q4 = list[13]
        comment4 = list[14]
        q5 = list[15]
        comment5 = list[16]
        q6 = list[17]
        comment6 = list[18]

        print('%s/%s - %s %s' % (str(i - 1), str(sheet.nrows - 2), firstName, lastName))
        # print(comments)
        if q1 != '':
            writeIndividualFile(doc_template_path, firstName, lastName, email, total, q1, q2, q3, q4, q5, q6,
                                comment1, comment2, comment3, comment4, comment5, comment6)
            # writeIndividualFile(doc_template_path, firstName, lastName, email, total, q1, q2, q3, q4,q5,
            #                     comment1, comment2, comment3, comment4, comment5)

